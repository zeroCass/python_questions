from tkinter.font import BOLD, ITALIC
import docx
import os

os.chdir(os.path.dirname(os.path.abspath(__file__)))
os.chdir('..\\automate_online_materials')

doc = docx.Document('demo.docx')
print(f'Numbers of paragraphs: {len(doc.paragraphs)}')

full_text = []

# paragraphs return a list of objs
for text in doc.paragraphs:
    # each obj has funtion text wich returns a string
    full_text.append(text.text)
print('\n'.join(full_text))


# each paragraph has a list of run
# run -> obj that contais texts with the same format
for idx, paragraph in enumerate(doc.paragraphs):
    print(f'\nParagraph[{idx}]_text: {paragraph.text}')
    for idx, run, in enumerate(paragraph.runs):
        print(f'run[{idx}]_text: {run.text}')



# is possible change the styles of the paragraphs and runs
# cehck the documentation for know more


# create a new docx file
doc_file = docx.Document()

# add titles with int 0-4
doc_file.add_heading('Titulo - 1', 0)
doc_file.add_heading('Titulo - 2', 4).add_run().add_break(docx.enum.text.WD_BREAK.LINE)

doc_file.add_paragraph('Hello, world!')

# add paragraph obj
doc_file.add_paragraph()

# adds objs runs into paragraph obj
doc_file.paragraphs[1].add_run('hello_')
doc_file.paragraphs[1].add_run('world.')

doc_file.save('helloworld.docx')

    







         


